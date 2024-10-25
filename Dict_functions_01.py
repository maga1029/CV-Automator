from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def fun_name(f_name_value, f_name_doc):
    paragraph = f_name_doc.add_paragraph()
    run = paragraph.add_run({f_name_value})
    run.bold = True
    font = run.font
    font.name = "Calibri"
    font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_two = f_name_doc.add_paragraph()
    run_two = paragraph_two.add_run("_" * 100)
    run_two.bold = True
    paragraph_two.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f"Name: {f_name_value}")


def fun_two(f_list_two, f_two_doc):
    string_two = ""
    for _ in range(len(f_list_two)):
        if f_list_two[_][1] == 2.1:
            string_two += f"{f_list_two[_][2]} • "
        if f_list_two[_][1] == 2.2:
            string_two += f"{f_list_two[_][2]} • "
        if f_list_two[_][1] == 2.3:
            string_two += f"{f_list_two[_][2]} • "
        if f_list_two[_][1] == 2.4:
            if f_list_two[_][1] == 2.4 and f_list_two[_] != f_list_two[-1]:
                string_two += f"{f_list_two[_][2]} • "
            else:
                string_two += f"{f_list_two[_][2]}"
    paragraph = f_two_doc.add_paragraph()
    run = paragraph.add_run(string_two)
    run.bold = False
    font = run.font
    font.name = "Calibri"
    font.size = Pt(11)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(string_two)


def fun_three(f_list_three, f_three_doc):
    if f_list_three:
        paragraph_one = f_three_doc.add_paragraph()
        run_one = paragraph_one.add_run("Personal Statement")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(len(f_list_three)):
        paragraph_two = f_three_doc.add_paragraph()
        run_two = paragraph_two.add_run(f_list_three[_][2])
        run_two.bold = False
        font_two = run_two.font
        font_two.name = "Calibri"
        font_two.size = Pt(11)
        paragraph_two.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def fun_four(f_list_four, f_four_doc):
    string_four_one = ""
    string_four_two = ""
    string_four_three = ""
    if f_list_four:
        paragraph_one = f_four_doc.add_paragraph()
        run_one = paragraph_one.add_run("Education")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_four)):
            if f_list_four[_][1] == 4.1:
                if string_four_one != "":
                    if string_four_one != "":
                        paragraph_two = f_four_doc.add_paragraph()
                        run_two = paragraph_two.add_run(string_four_one)
                        run_two.bold = False
                        font_two = run_two.font
                        font_two.name = "Calibri"
                        font_two.size = Pt(11)
                        paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if string_four_two != "":
                        paragraph_three = f_four_doc.add_paragraph()
                        run_three = paragraph_three.add_run(string_four_two)
                        run_three.bold = False
                        font_three = run_three.font
                        font_three.name = "Calibri"
                        font_three.size = Pt(11)
                        paragraph_three.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if string_four_three != "":
                        paragraph_four = f_four_doc.add_paragraph()
                        run_four = paragraph_four.add_run(string_four_three)
                        run_four.bold = False
                        font_four = run_four.font
                        font_four.name = "Calibri"
                        font_four.size = Pt(11)
                        paragraph_four.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    string_four_one = ""
                    string_four_two = ""
                    string_four_three = ""
                string_four_one += f_list_four[_][2]
            if f_list_four[_][1] == 4.2:
                if string_four_one == "":
                    string_four_one += f_list_four[_][2]
                else:
                    string_four_one += f", {f_list_four[_][2]}"
            if f_list_four[_][1] == 4.3:
                string_four_two += f"{f_list_four[_][2]}. "
            if f_list_four[_][1] == 4.4:
                string_four_two += f"{f_list_four[_][2]}. "
            if f_list_four[_][1] == 4.5:
                if string_four_three == "":
                    string_four_three += f"Relevant coursework: {f_list_four[_][2]}. "
                else:
                    string_four_three += f"{f_list_four[_][2]}. "

        if string_four_one != "":
            paragraph_two = f_four_doc.add_paragraph()
            run_two = paragraph_two.add_run(string_four_one)
            run_two.bold = False
            font_two = run_two.font
            font_two.name = "Calibri"
            font_two.size = Pt(11)
            paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if string_four_two != "":
            paragraph_three = f_four_doc.add_paragraph()
            run_three = paragraph_three.add_run(string_four_two)
            run_three.bold = False
            font_three = run_three.font
            font_three.name = "Calibri"
            font_three.size = Pt(11)
            paragraph_three.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if string_four_three != "":
            paragraph_four = f_four_doc.add_paragraph()
            run_four = paragraph_four.add_run(string_four_three)
            run_four.bold = False
            font_four = run_four.font
            font_four.name = "Calibri"
            font_four.size = Pt(11)
            paragraph_four.alignment = WD_ALIGN_PARAGRAPH.LEFT

        print(string_four_one)
        print(string_four_two)
        print(string_four_three)


def fun_five(f_list_five, f_five_doc):
    string_five_one = ""
    string_five_two = ""
    if f_list_five:
        paragraph_one = f_five_doc.add_paragraph()
        run_one = paragraph_one.add_run("Professional Experience")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_five)):
            if f_list_five[_][1] == 5.1:
                string_five_one += f"{f_list_five[_][2]}, "
            if f_list_five[_][1] == 5.2:
                string_five_one += f"{f_list_five[_][2]}, "
            if f_list_five[_][1] == 5.3:
                string_five_one += f"{f_list_five[_][2]}, "
            if f_list_five[_][1] == 5.4:
                try:
                    if f_list_five[_+1][1] == 5.5:
                        string_five_one += f"{f_list_five[_][2]} - "
                    else:
                        string_five_one += f"{f_list_five[_][2]}"
                except IndexError:
                    string_five_one += f"{f_list_five[_][2]}"
                    continue
            if f_list_five[_][1] == 5.5:
                string_five_one += f"{f_list_five[_][2]}"
            if f_list_five[_][1] == 5.6:
                if string_five_one != "":
                    print(string_five_one)
                    paragraph_two = f_five_doc.add_paragraph()
                    run_two = paragraph_two.add_run(string_five_one)
                    run_two.bold = True
                    font_two = run_two.font
                    font_two.name = "Calibri"
                    font_two.size = Pt(11)
                    paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    string_five_one = ""
                string_five_two += f_list_five[_][2]
                print(string_five_two)
                paragraph_three = f_five_doc.add_paragraph(style="ListBullet")
                run_three = paragraph_three.add_run(string_five_two)
                run_three.bold = False
                font_three = run_three.font
                font_three.name = "Calibri"
                font_three.size = Pt(11)
                paragraph_three.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                string_five_two = ""
        if string_five_one != "":
            print(string_five_one)
            paragraph_two = f_five_doc.add_paragraph()
            run_two = paragraph_two.add_run(string_five_one)
            run_two.bold = True
            font_two = run_two.font
            font_two.name = "Calibri"
            font_two.size = Pt(11)
            paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT


def fun_six(f_list_six, f_six_doc):
    print(f_list_six)
    if f_list_six:
        paragraph_one = f_six_doc.add_paragraph()
        run_one = paragraph_one.add_run("Publications and Presentations")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_six)):
            string_six_one = f_list_six[_][2]
            paragraph_two = f_six_doc.add_paragraph()
            run_two = paragraph_two.add_run(string_six_one)
            run_two.bold = False
            font_two = run_two.font
            font_two.name = "Calibri"
            font_two.size = Pt(11)
            paragraph_two.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            print(string_six_one)


def fun_seven(f_list_seven, f_seven_doc):
    string_seven_one = ""
    string_seven_two = ""
    if f_list_seven:
        paragraph_one = f_seven_doc.add_paragraph()
        run_one = paragraph_one.add_run("Certifications")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_seven)):
            if f_list_seven[_][1] == 7.1:
                if string_seven_one != "":
                    paragraph_two = f_seven_doc.add_paragraph()
                    run_two_one = paragraph_two.add_run(string_seven_one)
                    run_two_one.bold = True
                    font_two_one = run_two_one.font
                    font_two_one.name = "Calibri"
                    font_two_one.size = Pt(11)
                    run_two_two = paragraph_two.add_run(string_seven_two)
                    run_two_two.bold = False
                    font_two_two = run_two_two.font
                    font_two_two.name = "Calibri"
                    font_two_two.size = Pt(11)
                    paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    print(string_seven_one)
                    print(string_seven_two)
                    string_seven_two = ""
                string_seven_one = f"{f_list_seven[_][2]}: "
            if f_list_seven[_][1] == 7.2:
                try:
                    if f_list_seven[_+1][1] == 7.2:
                        string_seven_two += f"{f_list_seven[_][2]}, "
                    else:
                        string_seven_two += f"{f_list_seven[_][2]}."
                except IndexError:
                    string_seven_two += f"{f_list_seven[_][2]}."
                    continue
        paragraph_two = f_seven_doc.add_paragraph()
        run_two_one = paragraph_two.add_run(string_seven_one)
        run_two_one.bold = True
        font_two_one = run_two_one.font
        font_two_one.name = "Calibri"
        font_two_one.size = Pt(11)
        run_two_two = paragraph_two.add_run(string_seven_two)
        run_two_two.bold = False
        font_two_two = run_two_two.font
        font_two_two.name = "Calibri"
        font_two_two.size = Pt(11)
        paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
        print(string_seven_one)
        print(string_seven_two)


def fun_eight(f_list_eight, f_eight_doc):
    if f_list_eight:
        flag_1, flag_2, flag_3, flag_4, flag_5 = True, True, True, True, True
        string_8_1_2, string_8_2_2, string_8_3_2, string_8_4_2, string_8_5_2 = "", "", "", "", ""
        paragraph_one = f_eight_doc.add_paragraph()
        run_one = paragraph_one.add_run("Skills")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_eight)):
            if f_list_eight[_][1] == 8.1:
                try:
                    if f_list_eight[_+1][1] == 8.1:
                        string_8_1_2 += f"{f_list_eight[_][2]}, "
                    else:
                        string_8_1_2 += f"{f_list_eight[_][2]}."
                except IndexError:
                    string_8_1_2 += f"{f_list_eight[_][2]}."
                    continue
            if f_list_eight[_][1] == 8.2:
                try:
                    if f_list_eight[_+1][1] == 8.2:
                        string_8_2_2 += f"{f_list_eight[_][2]}, "
                    else:
                        string_8_2_2 += f"{f_list_eight[_][2]}."
                except IndexError:
                    string_8_2_2 += f"{f_list_eight[_][2]}."
                    continue
            if f_list_eight[_][1] == 8.3:
                try:
                    if f_list_eight[_+1][1] == 8.3:
                        string_8_3_2 += f"{f_list_eight[_][2]}, "
                    else:
                        string_8_3_2 += f"{f_list_eight[_][2]}."
                except IndexError:
                    string_8_3_2 += f"{f_list_eight[_][2]}."
                    continue
            if f_list_eight[_][1] == 8.4:
                try:
                    if f_list_eight[_+1][1] == 8.4:
                        string_8_4_2 += f"{f_list_eight[_][2]}, "
                    else:
                        string_8_4_2 += f"{f_list_eight[_][2]}."
                except IndexError:
                    string_8_4_2 += f"{f_list_eight[_][2]}."
                    continue
            if f_list_eight[_][1] == 8.5:
                try:
                    if f_list_eight[_+1][1] == 8.5:
                        string_8_5_2 += f"{f_list_eight[_][2]}, "
                    else:
                        string_8_5_2 += f"{f_list_eight[_][2]}."
                except IndexError:
                    string_8_5_2 += f"{f_list_eight[_][2]}."
                    continue
        print(string_8_1_2, string_8_2_2, string_8_3_2, string_8_4_2, string_8_5_2)
        for _ in range(len(f_list_eight)):
            if f_list_eight[_][1] == 8.1 and flag_1:
                string_8_1_1 = "Technical Tools: "
                paragraph_two = f_eight_doc.add_paragraph()
                run_2_1 = paragraph_two.add_run(string_8_1_1)
                run_2_1.bold = True
                font_2_1 = run_2_1.font
                font_2_1.name = "Calibri"
                font_2_1.size = Pt(11)
                run_2_2 = paragraph_two.add_run(string_8_1_2)
                run_2_2.bold = False
                font_2_2 = run_2_2.font
                font_2_2.name = "Calibri"
                font_2_2.size = Pt(11)
                paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
                flag_1 = False
            if f_list_eight[_][1] == 8.2 and flag_2:
                string_8_2_1 = "Programming Skills: "
                paragraph_three = f_eight_doc.add_paragraph()
                run_3_1 = paragraph_three.add_run(string_8_2_1)
                run_3_1.bold = True
                font_3_1 = run_3_1.font
                font_3_1.name = "Calibri"
                font_3_1.size = Pt(11)
                run_3_2 = paragraph_three.add_run(string_8_2_2)
                run_3_2.bold = False
                font_3_2 = run_3_2.font
                font_3_2.name = "Calibri"
                font_3_2.size = Pt(11)
                paragraph_three.alignment = WD_ALIGN_PARAGRAPH.LEFT
                flag_2 = False
            if f_list_eight[_][1] == 8.3 and flag_3:
                string_8_3_1 = "Languages: "
                paragraph_four = f_eight_doc.add_paragraph()
                run_4_1 = paragraph_four.add_run(string_8_3_1)
                run_4_1.bold = True
                font_4_1 = run_4_1.font
                font_4_1.name = "Calibri"
                font_4_1.size = Pt(11)
                run_4_2 = paragraph_four.add_run(string_8_3_2)
                run_4_2.bold = False
                font_4_2 = run_4_2.font
                font_4_2.name = "Calibri"
                font_4_2.size = Pt(11)
                paragraph_four.alignment = WD_ALIGN_PARAGRAPH.LEFT
                flag_3 = False
            if f_list_eight[_][1] == 8.4 and flag_4:
                string_8_4_1 = "Soft Skills: "
                paragraph_five = f_eight_doc.add_paragraph()
                run_5_1 = paragraph_five.add_run(string_8_4_1)
                run_5_1.bold = True
                font_5_1 = run_5_1.font
                font_5_1.name = "Calibri"
                font_5_1.size = Pt(11)
                run_5_2 = paragraph_five.add_run(string_8_4_2)
                run_5_2.bold = False
                font_5_2 = run_5_2.font
                font_5_2.name = "Calibri"
                font_5_2.size = Pt(11)
                paragraph_five.alignment = WD_ALIGN_PARAGRAPH.LEFT
                flag_4 = False
            if f_list_eight[_][1] == 8.5 and flag_5:
                string_8_5_1 = "Lab Skills: "
                paragraph_six = f_eight_doc.add_paragraph()
                run_6_1 = paragraph_six.add_run(string_8_5_1)
                run_6_1.bold = True
                font_6_1 = run_6_1.font
                font_6_1.name = "Calibri"
                font_6_1.size = Pt(11)
                run_6_2 = paragraph_six.add_run(string_8_5_2)
                run_6_2.bold = False
                font_6_2 = run_6_2.font
                font_6_2.name = "Calibri"
                font_6_2.size = Pt(11)
                paragraph_six.alignment = WD_ALIGN_PARAGRAPH.LEFT
                flag_5 = False


def fun_nine(f_list_nine, f_nine_doc):
    if f_list_nine:
        paragraph_one = f_nine_doc.add_paragraph()
        run_one = paragraph_one.add_run("Honors and Awards")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_nine)):
            string_nine_one = f_list_nine[_][2]
            paragraph_two = f_nine_doc.add_paragraph()
            run_two = paragraph_two.add_run(string_nine_one)
            run_two.bold = False
            font_two = run_two.font
            font_two.name = "Calibri"
            font_two.size = Pt(11)
            paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(string_nine_one)


def fun_ten(f_list_ten, f_ten_doc):
    string_10_2 = ""
    if f_list_ten:
        paragraph_one = f_ten_doc.add_paragraph()
        run_one = paragraph_one.add_run("Extracurricular Activities")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_ten)):
            if f_list_ten[_][1] == 10.1:
                string_ten_one = f_list_ten[_][2]
                paragraph_two = f_ten_doc.add_paragraph()
                run_two = paragraph_two.add_run(string_ten_one)
                run_two.bold = False
                font_two = run_two.font
                font_two.name = "Calibri"
                font_two.size = Pt(11)
                paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
                print(string_ten_one)
            if f_list_ten[_][1] == 10.2:
                try:
                    if f_list_ten[_+1][1] == 10.2:
                        string_10_2 += f"{f_list_ten[_][2]}, "
                    else:
                        string_10_2 += f"{f_list_ten[_][2]}."
                except IndexError:
                    string_10_2 += f"{f_list_ten[_][2]}."
                    continue
        print(string_10_2)
        if string_10_2:
            paragraph_three = f_ten_doc.add_paragraph()
            run_3_1 = paragraph_three.add_run("Hobbies: ")
            run_3_1.bold = True
            font_3_1 = run_3_1.font
            font_3_1.name = "Calibri"
            font_3_1.size = Pt(11)
            run_3_2 = paragraph_three.add_run(string_10_2)
            run_3_2.bold = False
            font_3_2 = run_3_2.font
            font_3_2.name = "Calibri"
            font_3_2.size = Pt(11)
            paragraph_three.alignment = WD_ALIGN_PARAGRAPH.LEFT


def fun_eleven(f_list_eleven, f_eleven_doc):
    if f_list_eleven:
        paragraph_one = f_eleven_doc.add_paragraph()
        run_one = paragraph_one.add_run("References")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_eleven)):
            string_eleven_one = f_list_eleven[_][2]
            paragraph_two = f_eleven_doc.add_paragraph()
            run_two = paragraph_two.add_run(string_eleven_one)
            run_two.bold = False
            font_two = run_two.font
            font_two.name = "Calibri"
            font_two.size = Pt(11)
            paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(string_eleven_one)


def fun_twelve(f_list_twelve, f_twelve_doc):
    if f_list_twelve:
        paragraph_one = f_twelve_doc.add_paragraph()
        run_one = paragraph_one.add_run("Professional Affiliations")
        run_one.bold = True
        font_one = run_one.font
        font_one.name = "Calibri"
        font_one.size = Pt(11)
        paragraph_one.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(len(f_list_twelve)):
            string_twelve_one = f_list_twelve[_][2]
            paragraph_two = f_twelve_doc.add_paragraph()
            run_two = paragraph_two.add_run(string_twelve_one)
            run_two.bold = False
            font_two = run_two.font
            font_two.name = "Calibri"
            font_two.size = Pt(11)
            paragraph_two.alignment = WD_ALIGN_PARAGRAPH.LEFT
            print(string_twelve_one)


if __name__ == "__main__":
    pass
