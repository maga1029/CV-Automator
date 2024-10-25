from functions import *
from docx import Document
from docx.shared import Cm
from docx2pdf import convert
from tkinter import messagebox
import pandas as pd

file_excel = "CV Excel.xlsx"
file_word = "test_cv.docx"
file_pdf = "test_cv.pdf"


def fun_read_excel(f1_file_excel):
    list_fields_cv = []
    df_excel = pd.read_excel(f1_file_excel).to_numpy()
    for _ in range(len(df_excel)):
        try:
            if pd.isna(df_excel[_][0]):
                df_excel[_][0] = ""
            if pd.isna(df_excel[_][3]):
                df_excel[_][3] = ""
            list_fields_cv.extend([[df_excel[_][2], df_excel[_][1], df_excel[_][3]]])
        except IndexError:
            messagebox.showinfo(title="Formating Error", message="Please format the Excel file according to "
                                                                 "instructions")
            return
    return list_fields_cv


def fun_write_word(f2_file_word, f2_list_fields, f2_file_pdf, no_space, to_pdf):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.26)
    section.bottom_margin = Cm(0.49)
    section.left_margin = Cm(1.06)
    section.right_margin = Cm(1.06)
    list_components_two = []
    list_components_three = []
    list_components_four = []
    list_components_five = []
    list_components_six = []
    list_components_seven = []
    list_components_eight = []
    list_components_nine = []
    list_components_ten = []
    list_components_eleven = []
    list_components_twelve = []

    for _ in range(len(f2_list_fields)):
        if f2_list_fields[_][0] == 1:
            if f2_list_fields[_][1] == 1.0:
                fun_name(f2_list_fields[_][2], doc)
            if f2_list_fields[_][1] == 2.1 or f2_list_fields[_][1] == 2.2 or f2_list_fields[_][1] == 2.3 or \
                    f2_list_fields[_][1] == 2.4:
                list_components_two.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 3.0:
                list_components_three.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 4.1 or f2_list_fields[_][1] == 4.2 or f2_list_fields[_][1] == 4.3 or \
                    f2_list_fields[_][1] == 4.4 or f2_list_fields[_][1] == 4.5:
                list_components_four.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 5.1 or f2_list_fields[_][1] == 5.2 or f2_list_fields[_][1] == 5.3 or \
                    f2_list_fields[_][1] == 5.4 or f2_list_fields[_][1] == 5.5 or f2_list_fields[_][1] == 5.6:
                list_components_five.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 6.0:
                list_components_six.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 7.1 or f2_list_fields[_][1] == 7.2:
                list_components_seven.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 8.1 or f2_list_fields[_][1] == 8.2 or f2_list_fields[_][1] == 8.3 or \
                    f2_list_fields[_][1] == 8.4 or f2_list_fields[_][1] == 8.5:
                list_components_eight.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 9.0:
                list_components_nine.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 10.1 or f2_list_fields[_][1] == 10.2:
                list_components_ten.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 11.0:
                list_components_eleven.append(f2_list_fields[_])
            if f2_list_fields[_][1] == 12.0:
                list_components_twelve.append(f2_list_fields[_])

    fun_two(list_components_two, doc)
    fun_three(list_components_three, doc)
    fun_four(list_components_four, doc)
    fun_five(list_components_five, doc)
    fun_six(list_components_six, doc)
    fun_seven(list_components_seven, doc)
    fun_eight(list_components_eight, doc)
    fun_nine(list_components_nine, doc)
    fun_ten(list_components_ten, doc)
    fun_eleven(list_components_eleven, doc)
    fun_twelve(list_components_twelve, doc)
    doc.save(f2_file_word)

    if no_space:
        doc = Document(f2_file_word)
        for _ in doc.paragraphs:
            _.paragraph_format.space_after = Pt(0)
        doc.save(f2_file_word)

    if to_pdf:
        convert(f2_file_word, f2_file_pdf)


if __name__ == "__main__":
    fun_write_word(file_word, fun_read_excel(file_excel), file_pdf, 1, 1)
    messagebox.showinfo(title="Finished", message="Process finished")
